"""
Dashboard for the Tenacious Conversion Agent.

GET  /                         – white-theme browser UI
POST /api/pipeline/run         – start async step-by-step pipeline
GET  /api/pipeline/history     – all finished runs, newest first
GET  /api/pipeline/{id}        – poll a single run's state
POST /api/pipeline/{id}/approve – approve (+ optionally edit) composed email
POST /api/pipeline/{id}/reject  – cancel the run
POST /api/email/compose        – compose draft only (no send)
POST /api/email/send           – send a pre-composed draft
POST /api/gap/analyze          – competitor gap for any lead email
POST /api/crm/sync             – manual HubSpot upsert
GET  /api/leads                – all leads (JSON)
GET  /api/leads/{email_b64}    – single lead detail
DELETE /api/leads/{email_b64}  – remove a lead
"""

import asyncio
import base64
import csv
import dataclasses as _dc
import hashlib
import hmac
import io
import logging
import os
import re
import time
import uuid
from concurrent.futures import ThreadPoolExecutor

from dotenv import load_dotenv
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, JSONResponse
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address

load_dotenv()

from agent import db
from agent import enrichment_pipeline as enrich_mod
from agent import email_outreach as email_mod
from agent import conversation_handler as conv
from agent import booking_handler as booking
from agent import hubspot_sync as hs
from agent import langfuse_logger as lf
from agent import sms_handler as sms_mod
from agent import competitor_gap as gap_mod
from agent import signals_research as signals_mod

_CALCOM_API_KEY        = os.getenv("CALCOM_API_KEY", "")
_AT_USERNAME           = os.getenv("AFRICA_TALKING_USERNAME", "sandbox")
_RESEND_WEBHOOK_SECRET = os.getenv("RESEND_WEBHOOK_SECRET", "")
_OUTBOUND_LIVE         = os.getenv("OUTBOUND_LIVE", "false").lower() in ("1", "true", "yes")
_STAFF_SINK            = os.getenv("STAFF_SINK_EMAIL", "sink@tenacious-pilot.dev")

logger = logging.getLogger(__name__)

limiter = Limiter(key_func=get_remote_address)
app = FastAPI(title="Tenacious Dashboard")
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

_executor      = ThreadPoolExecutor(max_workers=4)
_active_tasks: set          = set()
_RUN_STORE:    dict[str, dict] = {}
_RUN_HISTORY:  list[dict]      = []   # newest first, max 50

# Batch run state (single concurrent batch at a time)
_BATCH: dict = {
    "status":       "idle",   # idle|running|completed|failed
    "total":        0,
    "done":         0,
    "failed_count": 0,
    "current":      None,
    "results":      [],
    "started_at":   None,
    "completed_at": None,
}

# ── File parsing helpers ──────────────────────────────────────────────────────

_EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")

def _domain_from_url(url: str) -> str:
    url = url.strip().rstrip("/")
    url = re.sub(r"^https?://", "", url)
    url = re.sub(r"^www\.", "", url)
    return url.split("/")[0].split("?")[0]

def _parse_csv_bytes(content: bytes) -> list[dict]:
    text   = content.decode("utf-8", errors="replace")
    reader = csv.DictReader(io.StringIO(text))
    leads  = []
    for row in reader:
        r = {k.lower().strip(): (v or "").strip() for k, v in row.items()}
        email = ""
        for col in ("email", "contact_email", "email_address", "e-mail", "mail"):
            val = r.get(col, "")
            if val and "@" in val and "█" not in val and "■" not in val:
                email = val
                break
        domain = ""
        for col in ("website", "url", "domain", "site", "web", "homepage"):
            val = r.get(col, "")
            if val:
                domain = _domain_from_url(val)
                break
        if not email and domain:
            email = f"founder@{domain}"
        if not email:
            continue
        company = ""
        for col in ("name", "company", "company_name", "organization", "org", "account"):
            val = r.get(col, "")
            if val:
                company = val
                break
        leads.append({
            "email":   email,
            "company": company,
            "domain":  domain or email.split("@")[1],
            "source":  "direct" if ("founder@" not in email) else "synthesized",
        })
    return leads

def _emails_from_text(text: str) -> list[dict]:
    found  = list(dict.fromkeys(_EMAIL_RE.findall(text)))   # dedupe, preserve order
    return [
        {
            "email":   e,
            "company": e.split("@")[1].split(".")[0].title(),
            "domain":  e.split("@")[1],
            "source":  "extracted",
        }
        for e in found
    ]

def _parse_pdf_bytes(content: bytes) -> list[dict]:
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except ImportError:
        raise RuntimeError("pdfplumber not installed — run: pip install pdfplumber")
    return _emails_from_text(text)

def _parse_docx_bytes(content: bytes) -> list[dict]:
    try:
        from docx import Document
        doc  = Document(io.BytesIO(content))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text
    except ImportError:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")
    return _emails_from_text(text)


# ── Webhook signature helpers ─────────────────────────────────────────────────

def _verify_resend_signature(headers: dict, raw_body: bytes) -> bool:
    if not _RESEND_WEBHOOK_SECRET:
        return True
    msg_id     = headers.get("svix-id", "")
    timestamp  = headers.get("svix-timestamp", "")
    sig_header = headers.get("svix-signature", "")
    if not (msg_id and timestamp and sig_header):
        return False
    try:
        if abs(time.time() - int(timestamp)) > 300:
            return False
    except ValueError:
        return False
    signed       = f"{msg_id}.{timestamp}.{raw_body.decode()}"
    secret_bytes = base64.b64decode(_RESEND_WEBHOOK_SECRET.removeprefix("whsec_"))
    expected     = base64.b64encode(
        hmac.new(secret_bytes, signed.encode(), hashlib.sha256).digest()
    ).decode()
    for sig in sig_header.split(" "):
        _, actual = sig.split(",", 1) if "," in sig else ("v1", sig)
        if hmac.compare_digest(expected, actual):
            return True
    return False


def _verify_at_webhook(raw: dict) -> bool:
    return raw.get("username") == _AT_USERNAME


# ── Sync pipeline helpers (used by webhook routes) ────────────────────────────

def _run_full_pipeline(email: str, text: str) -> dict:
    trace_id        = lf.log_trace("first_contact", {"email": email, "text": text}, None)
    profile         = enrich_mod.enrich(email)
    lf.log_span(trace_id, "enrich", {"email": email}, profile.__dict__)
    company_signals = signals_mod.research(profile.domain, trace_id)
    profile.signals_research = company_signals.to_dict()
    lf.log_span(trace_id, "signals_research", {"domain": profile.domain}, profile.signals_research)
    gap_brief        = gap_mod.generate_competitor_gap_brief(profile, trace_id)
    send_result      = email_mod.compose_and_send(profile, trace_id)
    lead             = db.get_or_create(email)
    lead.status      = "outreach_sent"
    lead.profile     = {**_dc.asdict(profile), "competitor_gap_brief": gap_brief}
    name_parts       = email.split("@")[0].split(".")
    first            = name_parts[0].title()
    last             = name_parts[1].title() if len(name_parts) > 1 else ""
    contact_id       = hs.upsert_contact(
        email=email, first_name=first, last_name=last,
        company=profile.company_name,
        segment_label=email_mod.SEGMENT_LABELS[profile.segment],
        ai_maturity_score=profile.ai_maturity_score,
        booking_url="", enrichment_ts=profile.enriched_at, trace_id=trace_id,
    )
    lead.hubspot_contact_id = contact_id
    db.save_lead(lead)
    lf.log_trace("first_contact_complete", {"email": email}, send_result, session_id=lead.lead_id)
    return {
        "lead_id": lead.lead_id, "status": "outreach_sent",
        "send_result": send_result,
        "segment": email_mod.SEGMENT_LABELS[profile.segment],
        "ai_maturity_score": profile.ai_maturity_score,
        "competitor_gap_brief": gap_brief,
    }


def _run_reply_pipeline(identifier: str, text: str) -> dict:
    trace_id = lf.log_trace("reply", {"identifier": identifier, "text": text}, None)
    result   = conv.handle_reply(identifier, text, trace_id)
    lead     = db.get_or_create(identifier)
    if result["qualified"] and not lead.booking_url:
        name_parts     = identifier.split("@")[0].split(".") if "@" in identifier else [identifier]
        name           = " ".join(p.title() for p in name_parts)
        booking_result = booking.book(identifier, name, trace_id, api_key=_CALCOM_API_KEY)
        lead.booking_url = booking_result.get("booking_url", "")
        profile          = lead.profile
        hs.upsert_contact(
            email=identifier if "@" in identifier else profile.get("email", identifier),
            first_name=name_parts[0].title(),
            last_name=name_parts[1].title() if len(name_parts) > 1 else "",
            company=profile.get("company_name", ""),
            segment_label=email_mod.SEGMENT_LABELS.get(profile.get("segment", 0), "generic"),
            ai_maturity_score=profile.get("ai_maturity_score", 2),
            booking_url=lead.booking_url,
            enrichment_ts=profile.get("enriched_at", ""),
            trace_id=trace_id,
        )
        if lead.phone and lead.booking_url:
            sms_confirmation = sms_mod.send_booking_confirmation_sms(
                phone=lead.phone,
                booking_title=booking_result.get("title", "Discovery Call"),
                start_time=booking_result.get("start", ""),
                booking_url=lead.booking_url,
            )
            result["sms_confirmation"] = sms_confirmation
        result["booking_url"]    = lead.booking_url
        result["booking_result"] = booking_result
        db.save_lead(lead)
    lf.log_trace("reply_complete", {"identifier": identifier}, result, session_id=lead.lead_id)
    return result


# ── Async pipeline ────────────────────────────────────────────────────────────

def _push_to_history(run: dict) -> None:
    _RUN_HISTORY[:] = [r for r in _RUN_HISTORY if r["run_id"] != run["run_id"]]
    snap = {k: run.get(k) for k in (
        "run_id", "email", "status", "started_at", "completed_at",
        "result", "error", "gap_brief", "email_draft",
    )}
    snap["steps"]      = dict(run["steps"])
    snap["step_order"] = list(run["step_order"])
    _RUN_HISTORY.insert(0, snap)
    del _RUN_HISTORY[50:]


async def _pipeline_task(run_id: str) -> None:
    run   = _RUN_STORE[run_id]
    email = run["email"]
    text  = run["text"]
    loop  = asyncio.get_running_loop()

    def s_start(name: str) -> None:
        run["steps"][name] = {"name": name, "status": "running",
                              "data": None, "error": None, "ts": time.time()}
        if name not in run["step_order"]:
            run["step_order"].append(name)

    def s_done(name: str, data=None) -> None:
        if name in run["steps"]:
            run["steps"][name]["status"] = "done"
            run["steps"][name]["data"]   = data

    def s_err(name: str, err: str) -> None:
        if name in run["steps"]:
            run["steps"][name]["status"] = "error"
            run["steps"][name]["error"]  = err

    try:
        run["status"] = "running"

        # 1 ── Enrichment
        s_start("Enrichment")
        trace_id = lf.log_trace("first_contact", {"email": email, "text": text}, None)
        run["trace_id"] = trace_id
        profile = await loop.run_in_executor(_executor, enrich_mod.enrich, email)
        s_done("Enrichment", {
            "company":  profile.company_name,
            "domain":   profile.domain,
            "segment":  profile.segment,
            "cb_source":  (profile.crunchbase_signal or {}).get("source", "—"),
            "cb_conf":    round((profile.crunchbase_signal or {}).get("confidence", 0), 2),
            "jobs_roles": profile.open_engineering_roles,
            "layoffs":    profile.had_layoffs,
        })

        # 2 ── Signals Research
        s_start("Signals Research")
        company_signals = await loop.run_in_executor(
            _executor, signals_mod.research, profile.domain, trace_id
        )
        profile.signals_research = company_signals.to_dict()
        hook = (profile.signals_research or {}).get("personalization_hook", "")
        s_done("Signals Research", {"hook": (hook or "—")[:80]})

        # 3 ── Competitor Gap
        s_start("Competitor Gap")
        gap_brief = await loop.run_in_executor(
            _executor, gap_mod.generate_competitor_gap_brief, profile, trace_id
        )
        run["gap_brief"] = gap_brief
        s_done("Competitor Gap", {
            "angle":   gap_brief.get("recommended_angle", "")[:100],
            "summary": gap_brief.get("top_gap_summary", "")[:120],
        })

        # 4 ── Compose email — PAUSE for human approval
        s_start("Email Composition")
        subject, body_text = await loop.run_in_executor(
            _executor, email_mod.compose, profile, trace_id
        )
        det_ok, _ = email_mod._deterministic_tone_check(subject, body_text)
        if not det_ok:
            subject, body_text = await loop.run_in_executor(
                _executor, email_mod.compose, profile, trace_id
            )

        run["email_draft"] = {"to": email, "subject": subject, "body": body_text}
        run["profile"]     = _dc.asdict(profile)
        run["profile"]["competitor_gap_brief"] = gap_brief

        run["steps"]["Email Composition"]["status"] = "awaiting_approval"
        run["steps"]["Email Composition"]["data"]   = {
            "to": email, "subject": subject, "body": body_text,
        }
        run["status"] = "awaiting_approval"

        # wait up to 10 min for human to approve/reject
        for _ in range(1200):
            if run["status"] in ("approved", "rejected"):
                break
            await asyncio.sleep(0.5)

        if run["status"] == "rejected":
            run["steps"]["Email Composition"]["status"] = "rejected"
            run["status"] = "rejected"
            run["completed_at"] = time.time()
            _push_to_history(run)
            return

        if run["status"] != "approved":
            run["status"] = "failed"
            run["error"]  = "Email approval timed out (10 min)"
            run["completed_at"] = time.time()
            _push_to_history(run)
            return

        run["steps"]["Email Composition"]["status"] = "done"

        # 5 ── Send email
        draft = run["email_draft"]
        s_start("Email Send")
        send_result = await loop.run_in_executor(
            _executor, email_mod.send,
            draft["to"], draft["subject"], draft["body"], run["trace_id"],
        )
        s_done("Email Send", {
            "message_id": send_result.get("id", ""),
            "to": draft["to"],
        })

        # 6 ── HubSpot sync
        s_start("HubSpot Sync")
        lead        = db.get_or_create(email)
        lead.status = "outreach_sent"
        lead.profile = run["profile"]
        name_parts   = email.split("@")[0].split(".")
        first        = name_parts[0].title()
        last         = name_parts[1].title() if len(name_parts) > 1 else ""

        def _hs_sync():
            return hs.upsert_contact(
                email=email, first_name=first, last_name=last,
                company=profile.company_name,
                segment_label=email_mod.SEGMENT_LABELS[profile.segment],
                ai_maturity_score=profile.ai_maturity_score,
                booking_url="", enrichment_ts=profile.enriched_at,
                trace_id=run["trace_id"],
            )

        contact_id = await loop.run_in_executor(_executor, _hs_sync)
        lead.hubspot_contact_id = contact_id
        db.save_lead(lead)
        s_done("HubSpot Sync", {"contact_id": contact_id})

        run["status"] = "completed"
        run["result"] = {
            "lead_id":              lead.lead_id,
            "status":               "outreach_sent",
            "send_result":          send_result,
            "segment":              email_mod.SEGMENT_LABELS[profile.segment],
            "ai_maturity_score":    profile.ai_maturity_score,
            "competitor_gap_brief": gap_brief,
            "hubspot_contact_id":   contact_id,
        }
        lf.log_trace("first_contact_complete", {"email": email},
                     run["result"], session_id=lead.lead_id)

    except Exception as exc:
        logger.exception("Pipeline failed for %s", email)
        run["status"] = "failed"
        run["error"]  = str(exc)
        for s in run["steps"].values():
            if s["status"] == "running":
                s["status"] = "error"
                s["error"]  = str(exc)
    finally:
        run.setdefault("completed_at", time.time())
        _push_to_history(run)


# ── Health ────────────────────────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {"status": "ok", "ts": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())}


# ── Webhook routes ────────────────────────────────────────────────────────────

@app.post("/webhooks/email")
@limiter.limit("60/minute")
async def webhook_email(request: Request):
    raw_body = await request.body()
    if not _verify_resend_signature(dict(request.headers), raw_body):
        return JSONResponse({"error": "invalid signature"}, status_code=401)
    try:
        import json as _json
        body = _json.loads(raw_body)
    except Exception:
        return JSONResponse({"error": "invalid JSON"}, status_code=400)
    event_type = body.get("type", "")
    if event_type == "email.bounced":
        data        = body.get("data", {})
        to_list     = data.get("to", [])
        email       = to_list[0] if to_list else ""
        bounce_info = data.get("bounce", {})
        bounce_type = bounce_info.get("type", "soft")
        reason      = bounce_info.get("message", "")
        trace_id    = lf.log_trace("email_bounced", {"email": email, "type": bounce_type}, None)
        bounce_result = email_mod.handle_bounce(email, bounce_type, reason, trace_id)
        if email:
            hs.mark_bounced(email, bounce_type, trace_id)
            lead        = db.get_or_create(email)
            lead.status = "disqualified"
            db.save_lead(lead)
        return JSONResponse(bounce_result)
    if event_type == "email.complained":
        data    = body.get("data", {})
        to_list = data.get("to", [])
        email   = to_list[0] if to_list else ""
        trace_id = lf.log_trace("email_complaint", {"email": email}, None)
        complaint_result = email_mod.handle_complaint(email, trace_id)
        if email:
            hs.mark_bounced(email, "complaint", trace_id)
            lead        = db.get_or_create(email)
            lead.status = "disqualified"
            db.save_lead(lead)
        return JSONResponse(complaint_result)
    email     = body.get("from") or body.get("email", "")
    text      = body.get("text") or body.get("body", "")
    thread_id = body.get("thread_id", "")
    if not email:
        return JSONResponse({"error": "missing email"}, status_code=400)
    lead = db.get_or_create(email)
    if lead.status == "new" or not thread_id:
        result = _run_full_pipeline(email, text)
    else:
        result = _run_reply_pipeline(email, text)
    return JSONResponse(result)


@app.post("/webhooks/sms")
@limiter.limit("60/minute")
async def webhook_sms(request: Request):
    content_type = request.headers.get("content-type", "")
    if "application/x-www-form-urlencoded" in content_type or "multipart/form-data" in content_type:
        form = await request.form()
        raw  = dict(form)
    else:
        raw = await request.json()
    if not _verify_at_webhook(raw):
        return JSONResponse({"error": "unauthorized"}, status_code=401)
    parsed    = sms_mod.parse_at_payload(raw)
    phone     = parsed["phone"]
    text      = parsed["text"]
    if not phone:
        return JSONResponse({"error": "missing phone"}, status_code=400)
    trace_id  = lf.log_trace("sms_inbound", {"phone": phone, "text": text}, None)
    warm_lead = db.get_by_phone(phone)
    if warm_lead is None:
        return JSONResponse({
            "routed": False, "reason": "channel_hierarchy_gate",
            "detail": "No warm lead for this phone. First contact must be via email.",
        }, status_code=200)
    result = sms_mod.handle_inbound_sms(
        phone=phone, text=text, lead_status=warm_lead.status, trace_id=trace_id,
        reply_pipeline_fn=lambda _p, _t: _run_reply_pipeline(warm_lead.email, _t),
    )
    lf.log_trace("sms_complete", {"phone": phone}, result, session_id=warm_lead.lead_id)
    return JSONResponse(result)


@app.post("/simulate")
async def simulate(request: Request):
    body  = await request.json()
    email = body.get("email", "prospect@example.com")
    text  = body.get("text", "Tell me more about your engineering teams.")
    return JSONResponse(_run_full_pipeline(email, text))


@app.post("/simulate/sms")
async def simulate_sms(request: Request):
    body  = await request.json()
    email = body.get("email", "")
    phone = body.get("phone", "")
    text  = body.get("text", "Yes, let's talk.")
    if not email or not phone:
        return JSONResponse({"error": "email and phone required"}, status_code=400)
    db.link_phone(email, phone)
    raw      = {"from": phone, "text": text, "to": "Sandbox", "username": _AT_USERNAME}
    parsed   = sms_mod.parse_at_payload(raw)
    trace_id = lf.log_trace("sms_simulate", {"email": email, "phone": phone}, None)
    warm_lead = db.get_by_phone(phone)
    if warm_lead is None:
        return JSONResponse({"error": "lead not found after link_phone"}, status_code=500)
    result = sms_mod.handle_inbound_sms(
        phone=phone, text=parsed["text"], lead_status=warm_lead.status, trace_id=trace_id,
        reply_pipeline_fn=lambda _p, _t: _run_reply_pipeline(warm_lead.email, _t),
    )
    return JSONResponse(result)


# ── Pipeline API ──────────────────────────────────────────────────────────────

@app.get("/api/pipeline/history")
async def api_pipeline_history():
    return JSONResponse(_RUN_HISTORY)


@app.post("/api/pipeline/run")
async def api_pipeline_run(request: Request):
    body  = await request.json()
    email = body.get("email", "").strip()
    text  = body.get("text", "").strip()
    if not email:
        return JSONResponse({"error": "email required"}, status_code=400)
    run_id = str(uuid.uuid4())
    run = {
        "run_id": run_id, "email": email, "text": text,
        "status": "pending",
        "steps": {}, "step_order": [],
        "email_draft": None, "profile": None, "gap_brief": None,
        "trace_id": None, "result": None, "error": None,
        "started_at": time.time(), "completed_at": None,
    }
    _RUN_STORE[run_id] = run
    task = asyncio.create_task(_pipeline_task(run_id))
    _active_tasks.add(task)
    task.add_done_callback(_active_tasks.discard)
    return JSONResponse({"run_id": run_id})


@app.get("/api/pipeline/{run_id}")
async def api_pipeline_status(run_id: str):
    run = _RUN_STORE.get(run_id)
    if not run:
        return JSONResponse({"error": "not found"}, status_code=404)
    return JSONResponse({
        "run_id":       run["run_id"],
        "email":        run["email"],
        "status":       run["status"],
        "steps":        run["steps"],
        "step_order":   run["step_order"],
        "email_draft":  run["email_draft"],
        "result":       run.get("result"),
        "error":        run.get("error"),
        "started_at":   run["started_at"],
        "completed_at": run.get("completed_at"),
    })


@app.post("/api/pipeline/{run_id}/approve")
async def api_pipeline_approve(run_id: str, request: Request):
    run = _RUN_STORE.get(run_id)
    if not run:
        return JSONResponse({"error": "not found"}, status_code=404)
    if run["status"] != "awaiting_approval":
        return JSONResponse({"error": "not awaiting approval"}, status_code=400)
    body = await request.json()
    if run["email_draft"]:
        if "subject" in body:
            run["email_draft"]["subject"] = body["subject"]
        if "body" in body:
            run["email_draft"]["body"] = body["body"]
    run["status"] = "approved"
    return JSONResponse({"ok": True})


@app.post("/api/pipeline/{run_id}/reject")
async def api_pipeline_reject(run_id: str):
    run = _RUN_STORE.get(run_id)
    if not run:
        return JSONResponse({"error": "not found"}, status_code=404)
    run["status"] = "rejected"
    return JSONResponse({"ok": True})


# ── Email API ─────────────────────────────────────────────────────────────────

@app.post("/api/email/compose")
async def api_email_compose(request: Request):
    body  = await request.json()
    email = body.get("email", "").strip()
    if not email:
        return JSONResponse({"error": "email required"}, status_code=400)
    loop     = asyncio.get_running_loop()
    trace_id = lf.log_trace("compose_only", {"email": email}, None)
    profile  = await loop.run_in_executor(_executor, enrich_mod.enrich, email)
    subject, body_text = await loop.run_in_executor(_executor, email_mod.compose, profile, trace_id)
    return JSONResponse({
        "to":       email,
        "subject":  subject,
        "body":     body_text,
        "trace_id": trace_id,
        "company":  profile.company_name,
        "segment":  email_mod.SEGMENT_LABELS.get(profile.segment, "generic"),
    })


@app.post("/api/email/send")
async def api_email_send(request: Request):
    body      = await request.json()
    to        = body.get("to", "").strip()
    subject   = body.get("subject", "").strip()
    body_text = body.get("body", "").strip()
    trace_id  = body.get("trace_id", f"manual-{int(time.time())}")
    if not (to and subject and body_text):
        return JSONResponse({"error": "to, subject, body required"}, status_code=400)
    loop   = asyncio.get_running_loop()
    result = await loop.run_in_executor(_executor, email_mod.send, to, subject, body_text, trace_id)
    return JSONResponse(result)


# ── Gap API ───────────────────────────────────────────────────────────────────

@app.post("/api/gap/analyze")
async def api_gap_analyze(request: Request):
    body  = await request.json()
    email = body.get("email", "").strip()
    if not email:
        return JSONResponse({"error": "email required"}, status_code=400)
    loop             = asyncio.get_running_loop()
    trace_id         = lf.log_trace("gap_only", {"email": email}, None)
    profile          = await loop.run_in_executor(_executor, enrich_mod.enrich, email)
    company_signals  = await loop.run_in_executor(_executor, signals_mod.research, profile.domain, trace_id)
    profile.signals_research = company_signals.to_dict()
    gap_brief        = await loop.run_in_executor(
        _executor, gap_mod.generate_competitor_gap_brief, profile, trace_id
    )
    return JSONResponse({
        "gap_brief": gap_brief,
        "company":   profile.company_name,
        "domain":    profile.domain,
        "segment":   email_mod.SEGMENT_LABELS.get(profile.segment, "generic"),
    })


# ── CRM API ───────────────────────────────────────────────────────────────────

@app.post("/api/crm/sync")
async def api_crm_sync(request: Request):
    body  = await request.json()
    email = body.get("email", "").strip()
    if not email:
        return JSONResponse({"error": "email required"}, status_code=400)
    lead       = db.get_or_create(email)
    profile    = lead.profile or {}
    name_parts = email.split("@")[0].split(".")
    first      = name_parts[0].title()
    last       = name_parts[1].title() if len(name_parts) > 1 else ""
    loop       = asyncio.get_running_loop()
    trace_id   = lf.log_trace("crm_sync_manual", {"email": email}, None)

    def _sync():
        return hs.upsert_contact(
            email=email, first_name=first, last_name=last,
            company=profile.get("company_name", ""),
            segment_label=email_mod.SEGMENT_LABELS.get(profile.get("segment", 0), "generic"),
            ai_maturity_score=profile.get("ai_maturity_score", 0),
            booking_url=lead.booking_url or "",
            enrichment_ts=profile.get("enriched_at", ""),
            trace_id=trace_id,
        )

    contact_id = await loop.run_in_executor(_executor, _sync)
    lead.hubspot_contact_id = contact_id
    db.save_lead(lead)
    return JSONResponse({"ok": True, "contact_id": contact_id, "email": email})


# ── Batch API ─────────────────────────────────────────────────────────────────

@app.post("/api/batch/parse")
async def api_batch_parse(request: Request):
    form = await request.form()
    upload = form.get("file")
    if not upload:
        return JSONResponse({"error": "file required"}, status_code=400)
    filename = (upload.filename or "").lower()
    content  = await upload.read()
    try:
        if filename.endswith(".csv"):
            leads = _parse_csv_bytes(content)
        elif filename.endswith(".pdf"):
            leads = _parse_pdf_bytes(content)
        elif filename.endswith(".docx"):
            leads = _parse_docx_bytes(content)
        else:
            return JSONResponse({"error": "Unsupported file. Use .csv, .pdf or .docx"}, status_code=400)
    except Exception as exc:
        return JSONResponse({"error": str(exc)}, status_code=500)
    return JSONResponse({"count": len(leads), "leads": leads})


async def _batch_task(leads: list[dict]) -> None:
    global _BATCH
    _BATCH.update({
        "status": "running", "total": len(leads), "done": 0,
        "failed_count": 0, "current": None, "results": [],
        "started_at": time.time(), "completed_at": None,
    })
    loop = asyncio.get_running_loop()
    for lead in leads:
        email = lead["email"]
        _BATCH["current"] = email
        try:
            result = await loop.run_in_executor(
                _executor, _run_full_pipeline, email, ""
            )
            _BATCH["results"].append({"email": email, "status": "ok", "result": result})
            _BATCH["done"] += 1
        except Exception as exc:
            logger.exception("Batch pipeline failed for %s", email)
            _BATCH["results"].append({"email": email, "status": "error", "error": str(exc)})
            _BATCH["failed_count"] += 1
    _BATCH["status"]       = "completed"
    _BATCH["current"]      = None
    _BATCH["completed_at"] = time.time()


@app.post("/api/batch/run")
async def api_batch_run(request: Request):
    if _BATCH["status"] == "running":
        return JSONResponse({"error": "A batch is already running"}, status_code=409)
    body  = await request.json()
    leads = body.get("leads", [])
    n     = int(body.get("n", len(leads)))
    if not leads:
        return JSONResponse({"error": "leads list required"}, status_code=400)
    subset = leads[:n]
    task = asyncio.create_task(_batch_task(subset))
    _active_tasks.add(task)
    task.add_done_callback(_active_tasks.discard)
    return JSONResponse({"ok": True, "queued": len(subset)})


@app.get("/api/batch/status")
async def api_batch_status():
    return JSONResponse({k: v for k, v in _BATCH.items() if k != "results"} |
                        {"result_count": len(_BATCH["results"]),
                         "latest_results": _BATCH["results"][-10:]})


# ── Leads API ─────────────────────────────────────────────────────────────────

@app.get("/api/leads")
async def api_leads():
    leads = db.list_all()
    return JSONResponse([_lead_to_dict(l) for l in leads])


@app.get("/api/leads/{email_b64}")
async def api_lead_detail(email_b64: str):
    email = base64.urlsafe_b64decode(email_b64 + "==").decode()
    lead  = db.get_or_create(email)
    return JSONResponse(_lead_to_dict(lead))


@app.delete("/api/leads/{email_b64}")
async def api_delete_lead(email_b64: str):
    email   = base64.urlsafe_b64decode(email_b64 + "==").decode()
    deleted = db.delete_lead(email)
    return JSONResponse({"deleted": deleted, "email": email})


def _lead_to_dict(lead) -> dict:
    profile = lead.profile or {}
    return {
        "email":              lead.email,
        "lead_id":            lead.lead_id,
        "phone":              lead.phone,
        "status":             lead.status,
        "turns":              lead.turns,
        "created_at":         lead.created_at,
        "booking_url":        lead.booking_url,
        "hubspot_contact_id": lead.hubspot_contact_id,
        "company_name":       profile.get("company_name", ""),
        "domain":             profile.get("domain", ""),
        "segment":            profile.get("segment", 0),
        "segment_label":      email_mod.SEGMENT_LABELS.get(profile.get("segment", 0), "generic"),
        "ai_maturity_score":  profile.get("ai_maturity_score", 0),
        "enriched_at":        profile.get("enriched_at", ""),
        "history":            lead.history,
        "profile":            profile,
    }


# ── HTML ──────────────────────────────────────────────────────────────────────

_HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Tenacious Conversion Engine</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
<style>
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

:root {
  --bg:           #f1f5f9;
  --surface:      #ffffff;
  --surface-2:    #f8fafc;
  --border:       #e2e8f0;
  --border-2:     #cbd5e1;
  --text:         #0f172a;
  --text-2:       #334155;
  --muted:        #64748b;
  --accent:       #7c3aed;
  --accent-h:     #6d28d9;
  --accent-light: #ede9fe;
  --green:        #16a34a;
  --green-l:      #dcfce7;
  --yellow:       #d97706;
  --yellow-l:     #fef3c7;
  --red:          #dc2626;
  --red-l:        #fee2e2;
  --blue:         #2563eb;
  --blue-l:       #dbeafe;
  --radius:       10px;
  --topbar-h:     56px;
  --sidebar-w:    310px;
  --shadow:       0 1px 3px rgba(0,0,0,.08),0 1px 2px rgba(0,0,0,.05);
  --shadow-md:    0 4px 6px -1px rgba(0,0,0,.07),0 2px 4px -1px rgba(0,0,0,.05);
  --shadow-lg:    0 10px 15px -3px rgba(0,0,0,.08),0 4px 6px -2px rgba(0,0,0,.05);
}

body { font-family:'Inter',sans-serif; background:var(--bg); color:var(--text);
       min-height:100vh; font-size:14px; line-height:1.5; }

/* ── Topbar ── */
.topbar { height:var(--topbar-h); display:flex; align-items:center;
          justify-content:space-between; padding:0 24px;
          background:var(--surface); border-bottom:1px solid var(--border);
          position:sticky; top:0; z-index:200; box-shadow:var(--shadow); }
.logo { font-size:15px; font-weight:700; letter-spacing:-.3px; color:var(--text); }
.logo span { color:var(--accent); }
.topbar-right { display:flex; align-items:center; gap:14px; }
.health-pill { display:flex; align-items:center; gap:6px; padding:4px 10px;
               border-radius:99px; background:var(--surface-2);
               border:1px solid var(--border); font-size:12px; color:var(--muted); }
.hdot { width:7px; height:7px; border-radius:50%; background:#94a3b8;
        transition:background .3s; }
.hdot.ok { background:var(--green); box-shadow:0 0 5px rgba(22,163,74,.4); }
.outbound-badge { font-size:11px; padding:3px 8px; border-radius:99px; font-weight:600; }
.outbound-live { background:var(--green-l); color:var(--green); }
.outbound-sink { background:var(--yellow-l); color:var(--yellow); }

/* ── Layout ── */
.layout { display:flex; min-height:calc(100vh - var(--topbar-h)); align-items:flex-start; }
.main-area { flex:1; min-width:0; padding:24px; }
.sidebar { width:var(--sidebar-w); flex-shrink:0; border-left:1px solid var(--border);
           background:var(--surface); padding:16px;
           position:sticky; top:var(--topbar-h);
           height:calc(100vh - var(--topbar-h)); overflow-y:auto; }

/* ── Stats ── */
.stats { display:grid; grid-template-columns:repeat(auto-fit,minmax(120px,1fr));
         gap:10px; margin-bottom:20px; }
.stat-card { background:var(--surface); border:1px solid var(--border);
             border-radius:var(--radius); padding:14px 16px; box-shadow:var(--shadow); }
.stat-num { font-size:26px; font-weight:700; line-height:1; }
.stat-lbl { font-size:11px; color:var(--muted); text-transform:uppercase;
            letter-spacing:.6px; font-weight:500; margin-top:3px; }
.c-all  { color:var(--text); }
.c-new  { color:var(--muted); }
.c-sent { color:var(--blue); }
.c-conv { color:var(--yellow); }
.c-qual { color:var(--green); }
.c-disq { color:var(--red); }

/* ── Panel ── */
.panel { background:var(--surface); border:1px solid var(--border);
         border-radius:var(--radius); margin-bottom:16px;
         box-shadow:var(--shadow); overflow:hidden; }
.panel-head { display:flex; align-items:center; justify-content:space-between;
              padding:13px 18px; cursor:pointer; user-select:none;
              border-bottom:1px solid transparent; }
.panel-head:hover { background:var(--surface-2); }
.panel-head.open { border-bottom-color:var(--border); }
.panel-title { font-size:13px; font-weight:600; color:var(--text);
               display:flex; align-items:center; gap:8px; }
.panel-icon { font-size:15px; }
.chevron { font-size:10px; color:var(--muted); transition:transform .2s; }
.panel-head.open .chevron { transform:rotate(180deg); }
.panel-body { padding:18px; display:none; }
.panel-body.open { display:block; }

/* ── Forms ── */
.form-row { display:flex; gap:10px; flex-wrap:wrap; align-items:flex-end; margin-bottom:12px; }
.field { display:flex; flex-direction:column; gap:5px; flex:1; min-width:180px; }
label { font-size:11px; font-weight:600; color:var(--muted);
        text-transform:uppercase; letter-spacing:.5px; }
input, textarea, select {
  background:var(--surface-2); border:1px solid var(--border-2);
  color:var(--text); border-radius:7px; padding:8px 12px;
  font-family:inherit; font-size:13px; outline:none;
  transition:border-color .15s, box-shadow .15s; width:100%; }
input:focus, textarea:focus, select:focus {
  border-color:var(--accent); box-shadow:0 0 0 3px rgba(124,58,237,.1); }
textarea { resize:vertical; min-height:80px; }
select { cursor:pointer; }

.btn { padding:8px 16px; border-radius:7px; border:none; cursor:pointer;
       font-family:inherit; font-size:13px; font-weight:600;
       transition:opacity .15s,transform .1s,box-shadow .15s; white-space:nowrap;
       display:inline-flex; align-items:center; gap:6px; }
.btn:active { transform:scale(.97); }
.btn:disabled { opacity:.45; cursor:not-allowed; }
.btn-primary { background:var(--accent); color:#fff; box-shadow:0 1px 3px rgba(124,58,237,.3); }
.btn-primary:hover:not(:disabled) { background:var(--accent-h); box-shadow:0 2px 6px rgba(124,58,237,.4); }
.btn-success { background:var(--green); color:#fff; }
.btn-success:hover:not(:disabled) { opacity:.88; }
.btn-danger  { background:var(--red-l); color:var(--red); border:1px solid rgba(220,38,38,.25); }
.btn-danger:hover:not(:disabled)  { background:rgba(220,38,38,.2); }
.btn-ghost   { background:transparent; color:var(--muted);
               border:1px solid var(--border-2); }
.btn-ghost:hover { color:var(--text); border-color:var(--border-2); background:var(--surface-2); }
.btn-sm { padding:5px 10px; font-size:12px; }

/* ── Spinner ── */
.spin { display:inline-block; animation:spin .7s linear infinite; }
@keyframes spin { to { transform:rotate(360deg); } }
.btn-spinner { width:13px; height:13px; border:2px solid rgba(255,255,255,.35);
               border-top-color:#fff; border-radius:50%; animation:spin .7s linear infinite; }

/* ── Result box ── */
.result-box { background:var(--surface-2); border:1px solid var(--border);
              border-radius:7px; padding:14px; font-family:'Menlo','Consolas',monospace;
              font-size:12px; color:var(--text-2); white-space:pre-wrap; word-break:break-all;
              max-height:260px; overflow-y:auto; margin-top:12px; display:none; }
.result-box.visible { display:block; }
.result-box.error { border-color:rgba(220,38,38,.35); color:var(--red); }

/* ── Pipeline steps ── */
.pipe-form { margin-bottom:16px; }
.progress-wrap { margin:14px 0 6px; }
.progress-label { display:flex; justify-content:space-between;
                  font-size:11px; color:var(--muted); margin-bottom:5px; font-weight:500; }
.progress-track { width:100%; height:7px; background:var(--border);
                  border-radius:99px; overflow:hidden; }
.progress-fill  { height:100%; background:linear-gradient(90deg,var(--accent),#a855f7);
                  border-radius:99px; transition:width .6s cubic-bezier(.4,0,.2,1); width:0%; }
.steps-list { display:flex; flex-direction:column; gap:6px; }
.step-card { display:flex; align-items:center; gap:12px; padding:10px 14px;
             border-radius:8px; border:1px solid var(--border);
             background:var(--surface-2); transition:all .25s; }
.step-card.s-running  { border-color:#93c5fd; background:var(--blue-l); }
.step-card.s-done     { border-color:#86efac; background:var(--green-l); }
.step-card.s-error    { border-color:#fca5a5; background:var(--red-l); }
.step-card.s-awaiting { border-color:#fcd34d; background:var(--yellow-l); }
.step-card.s-rejected { border-color:#fca5a5; background:var(--red-l); opacity:.7; }
.step-card.s-pending  { opacity:.5; }
.step-icon-wrap { width:26px; height:26px; border-radius:50%; display:flex;
                  align-items:center; justify-content:center; font-size:13px;
                  font-weight:700; flex-shrink:0; }
.step-icon-wrap.s-pending  { background:#e2e8f0; color:#94a3b8; }
.step-icon-wrap.s-running  { background:var(--blue); color:#fff; }
.step-icon-wrap.s-done     { background:var(--green); color:#fff; }
.step-icon-wrap.s-error    { background:var(--red); color:#fff; }
.step-icon-wrap.s-awaiting { background:var(--yellow); color:#fff; }
.step-icon-wrap.s-rejected { background:var(--red); color:#fff; }
.step-content { flex:1; min-width:0; }
.step-name { font-size:13px; font-weight:600; color:var(--text); }
.step-detail { font-size:11px; color:var(--muted); margin-top:2px;
               overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
.step-badge { font-size:10px; font-weight:700; padding:2px 8px; border-radius:99px;
              text-transform:uppercase; letter-spacing:.5px; flex-shrink:0; }
.step-badge.s-pending  { background:#e2e8f0; color:#94a3b8; }
.step-badge.s-running  { background:var(--blue-l); color:var(--blue); }
.step-badge.s-done     { background:var(--green-l); color:var(--green); }
.step-badge.s-error    { background:var(--red-l); color:var(--red); }
.step-badge.s-awaiting { background:var(--yellow-l); color:var(--yellow); }
.step-badge.s-rejected { background:var(--red-l); color:var(--red); }
.pipe-status-msg { font-size:13px; color:var(--muted); margin-top:12px;
                   font-weight:500; text-align:center; }

/* ── Table ── */
.tbl-wrap { overflow-x:auto; }
table { width:100%; border-collapse:collapse; }
th { font-size:11px; font-weight:600; color:var(--muted); text-transform:uppercase;
     letter-spacing:.5px; padding:10px 14px; text-align:left;
     border-bottom:1px solid var(--border); white-space:nowrap;
     background:var(--surface-2); }
td { padding:11px 14px; border-bottom:1px solid var(--border);
     vertical-align:middle; font-size:13px; }
tr:last-child td { border-bottom:none; }
tbody tr { cursor:pointer; transition:background .1s; }
tbody tr:hover { background:var(--surface-2); }
.email-cell { font-family:'Menlo','Consolas',monospace; font-size:12px; color:var(--accent); }
.company-cell { font-weight:500; }
.empty-state { text-align:center; padding:40px; color:var(--muted); font-size:13px; }

/* ── Badges ── */
.badge { display:inline-flex; align-items:center; padding:3px 9px; border-radius:99px;
         font-size:11px; font-weight:600; white-space:nowrap; }
.badge-new  { background:#f1f5f9; color:#64748b; }
.badge-sent { background:var(--blue-l); color:var(--blue); }
.badge-conv { background:var(--yellow-l); color:var(--yellow); }
.badge-qual { background:var(--green-l); color:var(--green); }
.badge-disq { background:var(--red-l); color:var(--red); }
.score { display:inline-flex; align-items:center; gap:4px; }
.score-dot { width:7px; height:7px; border-radius:50%; }
.s0{background:#94a3b8} .s1{background:#22c55e} .s2{background:#f59e0b}
.s3{background:#f97316} .s4{background:#ef4444}

/* ── Toolbar ── */
.toolbar { display:flex; gap:10px; align-items:center; margin-bottom:14px; flex-wrap:wrap; }
.search-input { flex:1; min-width:180px; }

/* ── Sidebar ── */
.sidebar-title { font-size:12px; font-weight:700; text-transform:uppercase;
                 letter-spacing:.6px; color:var(--muted); margin-bottom:12px;
                 padding-bottom:10px; border-bottom:1px solid var(--border); }
.run-card { border:1px solid var(--border); border-radius:8px; margin-bottom:8px;
            background:var(--surface-2); cursor:pointer; overflow:hidden;
            transition:box-shadow .15s; }
.run-card:hover { box-shadow:var(--shadow-md); }
.run-card-head { padding:10px 12px; }
.run-email { font-size:12px; font-family:'Menlo','Consolas',monospace;
             color:var(--accent); font-weight:500; margin-bottom:3px;
             overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
.run-meta { display:flex; align-items:center; justify-content:space-between; }
.run-time { font-size:11px; color:var(--muted); }
.run-status-pill { font-size:10px; font-weight:700; padding:2px 7px;
                   border-radius:99px; text-transform:uppercase; letter-spacing:.4px; }
.rp-completed { background:var(--green-l); color:var(--green); }
.rp-failed    { background:var(--red-l); color:var(--red); }
.rp-rejected  { background:#f1f5f9; color:var(--muted); }
.rp-running   { background:var(--blue-l); color:var(--blue); }
.rp-awaiting  { background:var(--yellow-l); color:var(--yellow); }
.rp-pending   { background:#f1f5f9; color:var(--muted); }
.run-detail   { padding:0 12px 10px; display:none; border-top:1px solid var(--border);
                margin-top:2px; padding-top:8px; }
.run-detail.open { display:block; }
.run-result-pre { font-family:'Menlo','Consolas',monospace; font-size:10px;
                  color:var(--text-2); white-space:pre-wrap; word-break:break-all;
                  max-height:140px; overflow-y:auto; }
.run-gap-preview { font-size:11px; color:var(--muted); margin-top:6px;
                   font-style:italic; }
.run-error-msg { font-size:11px; color:var(--red); margin-top:4px; }
.sidebar-empty { text-align:center; padding:32px 8px; color:var(--muted); font-size:12px; }

/* ── Overlays / Modals ── */
.overlay { position:fixed; inset:0; background:rgba(15,23,42,.5); z-index:300;
           display:none; align-items:center; justify-content:center; padding:20px; }
.overlay.open { display:flex; }
.modal { background:var(--surface); border:1px solid var(--border-2);
         border-radius:14px; width:100%; max-width:640px; max-height:90vh;
         overflow-y:auto; box-shadow:var(--shadow-lg);
         animation:fadeUp .2s ease; }
.modal-lg { max-width:740px; }
@keyframes fadeUp { from{opacity:0;transform:translateY(10px)} to{opacity:1;transform:translateY(0)} }
.modal-head { padding:18px 22px 14px; border-bottom:1px solid var(--border);
              display:flex; align-items:flex-start; justify-content:space-between; }
.modal-head h3 { font-size:15px; font-weight:700; }
.modal-head .sub { font-size:12px; color:var(--muted); margin-top:2px; }
.close-btn { background:none; border:none; color:var(--muted); font-size:18px;
             cursor:pointer; padding:0 4px; line-height:1; border-radius:4px; }
.close-btn:hover { color:var(--text); background:var(--surface-2); }
.modal-body { padding:20px 22px; }
.modal-footer { display:flex; gap:10px; justify-content:flex-end;
                padding:14px 22px; border-top:1px solid var(--border);
                background:var(--surface-2); border-radius:0 0 14px 14px; }

/* ── Email approval modal specifics ── */
.approval-to-row { display:flex; align-items:center; gap:8px; padding:10px 12px;
                   background:var(--surface-2); border:1px solid var(--border);
                   border-radius:7px; margin-bottom:12px; }
.approval-to-label { font-size:11px; font-weight:700; color:var(--muted);
                     text-transform:uppercase; letter-spacing:.5px; white-space:nowrap; }
.approval-to-email { font-size:13px; font-family:'Menlo','Consolas',monospace;
                     color:var(--accent); font-weight:500; }
.sink-note { font-size:11px; color:var(--yellow); font-weight:500;
             margin-left:auto; background:var(--yellow-l); padding:2px 8px;
             border-radius:99px; white-space:nowrap; }
.approval-hint { font-size:12px; color:var(--muted); margin-bottom:16px;
                 padding:8px 12px; background:var(--accent-light);
                 border-radius:7px; border-left:3px solid var(--accent); }

/* ── Lead detail modal ── */
.kv-grid { display:grid; grid-template-columns:1fr 1fr; gap:12px; margin-bottom:16px; }
.kv .k  { font-size:11px; color:var(--muted); text-transform:uppercase;
           letter-spacing:.5px; margin-bottom:2px; font-weight:600; }
.kv .v  { font-size:13px; font-weight:500; word-break:break-all; }
.section-label { font-size:11px; font-weight:700; color:var(--muted); text-transform:uppercase;
                 letter-spacing:.6px; margin:16px 0 8px; border-top:1px solid var(--border);
                 padding-top:14px; }
.history-item { background:var(--surface-2); border:1px solid var(--border);
                border-radius:7px; padding:10px 12px; margin-bottom:6px; }
.history-item .role { font-size:10px; font-weight:700; text-transform:uppercase;
                      letter-spacing:.5px; margin-bottom:3px; }
.role-user      { color:var(--blue); }
.role-assistant { color:var(--accent); }
.history-item .msg { font-size:12px; color:var(--text-2); line-height:1.5; }
.json-block { background:var(--surface-2); border:1px solid var(--border); border-radius:7px;
              padding:10px; font-family:monospace; font-size:11px; color:var(--muted);
              white-space:pre-wrap; word-break:break-all; max-height:160px; overflow-y:auto; }

/* ── Info block (Resend guide) ── */
.info-block { background:var(--surface-2); border:1px solid var(--border);
              border-radius:8px; padding:14px 16px; font-size:13px; color:var(--text-2);
              line-height:1.7; }
.info-block h4 { font-size:13px; font-weight:700; color:var(--text); margin-bottom:8px; }
.info-block code { background:var(--accent-light); color:var(--accent); padding:1px 5px;
                   border-radius:4px; font-family:'Menlo','Consolas',monospace; font-size:12px; }
.info-block ol, .info-block ul { padding-left:20px; }
.info-block li { margin-bottom:4px; }

/* ── Batch upload ── */
.upload-zone { border:2px dashed var(--border-2); border-radius:10px;
               padding:32px 20px; text-align:center; cursor:pointer;
               transition:border-color .2s,background .2s; background:var(--surface-2); }
.upload-zone:hover,.upload-zone.drag-over { border-color:var(--accent);
               background:var(--accent-light); }
.upload-icon { font-size:32px; margin-bottom:8px; }
.upload-label { font-size:14px; font-weight:600; color:var(--text-2); margin-bottom:4px; }
.upload-hint  { font-size:12px; color:var(--muted); }
.batch-controls { display:flex; align-items:center; gap:14px; flex-wrap:wrap;
                  margin-bottom:14px; padding:12px 14px; background:var(--surface-2);
                  border:1px solid var(--border); border-radius:8px; }
.batch-controls .field { min-width:120px; flex:0 0 auto; }
.batch-count-badge { font-size:13px; font-weight:600; color:var(--accent);
                     background:var(--accent-light); padding:4px 12px;
                     border-radius:99px; white-space:nowrap; }
.checkbox-row { display:flex; align-items:center; gap:6px; font-size:13px;
                color:var(--text-2); cursor:pointer; user-select:none; }
.checkbox-row input[type=checkbox] { width:15px; height:15px; accent-color:var(--accent); }
.batch-table th, .batch-table td { padding:8px 12px; }
.batch-table .num-cell { color:var(--muted); font-size:12px; text-align:center; width:40px; }
.batch-progress-wrap { margin-top:16px; }
.batch-item { display:flex; align-items:center; gap:10px; padding:7px 12px;
              border-radius:7px; border:1px solid var(--border);
              background:var(--surface-2); margin-bottom:5px; font-size:13px; }
.batch-item .bi-email { font-family:'Menlo','Consolas',monospace; font-size:12px;
                         color:var(--accent); flex:1; min-width:0;
                         overflow:hidden; text-overflow:ellipsis; white-space:nowrap; }
.batch-item .bi-status { font-size:11px; font-weight:700; padding:2px 8px;
                          border-radius:99px; white-space:nowrap; }
.bi-ok      { background:var(--green-l); color:var(--green); }
.bi-error   { background:var(--red-l);   color:var(--red); }
.bi-running { background:var(--blue-l);  color:var(--blue); }
.bi-pending { background:#f1f5f9;        color:var(--muted); }

/* ── Scrollbar ── */
::-webkit-scrollbar { width:5px; height:5px; }
::-webkit-scrollbar-track { background:transparent; }
::-webkit-scrollbar-thumb { background:var(--border-2); border-radius:3px; }

/* ── Approval Alert Banner ── */
#approval-banner {
  display:none; position:fixed; top:0; left:0; right:0; z-index:500;
  background:#d97706; color:#fff; text-align:center; padding:13px 20px;
  font-weight:600; font-size:14px; cursor:pointer; letter-spacing:.01em;
  animation:pulse-banner 1.8s ease-in-out infinite;
}
#approval-banner.show { display:block; }
@keyframes pulse-banner {
  0%,100% { background:#d97706; }
  50%      { background:#b45309; }
}
</style>
</head>
<body>
<!-- __SERVER_CONFIG__ -->

<div id="approval-banner" onclick="document.getElementById('approval-overlay').classList.add('open')">
  ⚠️ Email ready — click here to review and approve before it sends
</div>

<div class="topbar">
  <div class="logo">Tenacious <span>Engine</span></div>
  <div class="topbar-right">
    <span id="outbound-badge" class="outbound-badge outbound-sink">SINK MODE</span>
    <div class="health-pill">
      <span class="hdot" id="hDot"></span>
      <span id="hLabel">checking…</span>
    </div>
    <button class="btn btn-ghost btn-sm" onclick="loadLeads()">↻ Refresh</button>
  </div>
</div>

<div class="layout">

<!-- ══════════════════ MAIN AREA ══════════════════ -->
<main class="main-area">

  <!-- Stats -->
  <div class="stats" id="stats">
    <div class="stat-card"><div class="stat-num c-all" id="s-total">—</div><div class="stat-lbl">Total</div></div>
    <div class="stat-card"><div class="stat-num c-new"  id="s-new">—</div><div class="stat-lbl">New</div></div>
    <div class="stat-card"><div class="stat-num c-sent" id="s-sent">—</div><div class="stat-lbl">Outreach Sent</div></div>
    <div class="stat-card"><div class="stat-num c-conv" id="s-conv">—</div><div class="stat-lbl">Conversing</div></div>
    <div class="stat-card"><div class="stat-num c-qual" id="s-qual">—</div><div class="stat-lbl">Qualified</div></div>
    <div class="stat-card"><div class="stat-num c-disq" id="s-disq">—</div><div class="stat-lbl">Disqualified</div></div>
  </div>

  <!-- ── Pipeline Runner ── -->
  <div class="panel">
    <div class="panel-head open" id="pipe-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">⚡</span>Pipeline Runner</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body open" id="pipe-body">
      <div class="pipe-form">
        <div class="form-row">
          <div class="field">
            <label>Lead Email</label>
            <input type="email" id="pipe-email" placeholder="prospect@company.com">
          </div>
          <div class="field">
            <label>Opening Message (optional)</label>
            <input type="text" id="pipe-text" placeholder="Tell me more about your engineering team.">
          </div>
          <button class="btn btn-primary" id="pipe-btn" onclick="startPipeline()">
            <span id="pipe-spin" style="display:none" class="btn-spinner"></span>
            Run Pipeline
          </button>
        </div>
      </div>

      <div id="pipe-progress-wrap" style="display:none">
        <div class="progress-wrap">
          <div class="progress-label">
            <span id="pipe-status-text">Starting…</span>
            <span id="pipe-pct">0%</span>
          </div>
          <div class="progress-track"><div class="progress-fill" id="pipe-fill"></div></div>
        </div>
        <div class="steps-list" id="pipe-steps"></div>
        <div class="pipe-status-msg" id="pipe-msg"></div>
      </div>
    </div>
  </div>

  <!-- ── Email Outreach ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">✉️</span>Email Outreach</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <p style="font-size:12px;color:var(--muted);margin-bottom:14px">
        Compose a draft email for any lead (runs enrichment + LLM compose). Review and edit before sending.
      </p>
      <div class="form-row">
        <div class="field">
          <label>Lead Email</label>
          <input type="email" id="em-email" placeholder="prospect@company.com">
        </div>
        <button class="btn btn-ghost" id="em-compose-btn" onclick="composeEmail()">
          <span id="em-compose-spin" style="display:none" class="btn-spinner" style="border-top-color:var(--accent)"></span>
          Compose Draft
        </button>
      </div>
      <div id="em-draft-wrap" style="display:none">
        <div class="approval-to-row" style="margin-bottom:12px">
          <span class="approval-to-label">To</span>
          <span class="approval-to-email" id="em-to"></span>
          <span id="em-sink-note" class="sink-note" style="display:none">→ redirected to sink</span>
        </div>
        <div class="field" style="margin-bottom:10px">
          <label>Subject</label>
          <input type="text" id="em-subject">
        </div>
        <div class="field" style="margin-bottom:14px">
          <label>Body</label>
          <textarea id="em-body" style="min-height:140px"></textarea>
        </div>
        <div style="display:flex;gap:10px">
          <button class="btn btn-success" onclick="sendComposedEmail()">
            <span id="em-send-spin" style="display:none" class="btn-spinner"></span>
            Approve &amp; Send
          </button>
          <button class="btn btn-ghost" onclick="document.getElementById('em-draft-wrap').style.display='none'">Cancel</button>
        </div>
      </div>
      <pre class="result-box" id="em-result"></pre>
    </div>
  </div>

  <!-- ── Competitor Gap Analysis ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">🎯</span>Competitor Gap Analysis</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <p style="font-size:12px;color:var(--muted);margin-bottom:14px">
        Runs enrichment + signals research + gap brief for any lead email.
      </p>
      <div class="form-row">
        <div class="field">
          <label>Lead Email</label>
          <input type="email" id="gap-email" placeholder="prospect@company.com">
        </div>
        <button class="btn btn-primary" id="gap-btn" onclick="analyzeGap()">
          <span id="gap-spin" style="display:none" class="btn-spinner"></span>
          Analyze Gap
        </button>
      </div>
      <pre class="result-box" id="gap-result"></pre>
    </div>
  </div>

  <!-- ── CRM Sync ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">🔗</span>CRM Sync (HubSpot)</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <p style="font-size:12px;color:var(--muted);margin-bottom:14px">
        Manually upsert a lead into HubSpot using their current profile from the DB.
      </p>
      <div class="form-row">
        <div class="field">
          <label>Lead Email</label>
          <input type="email" id="crm-email" placeholder="prospect@company.com">
        </div>
        <button class="btn btn-primary" id="crm-btn" onclick="syncCRM()">
          <span id="crm-spin" style="display:none" class="btn-spinner"></span>
          Sync to HubSpot
        </button>
      </div>
      <pre class="result-box" id="crm-result"></pre>
    </div>
  </div>

  <!-- ── SMS Simulation ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">📱</span>SMS Simulation</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <div class="form-row">
        <div class="field">
          <label>Lead Email</label>
          <input type="email" id="sms-email" placeholder="prospect@company.com">
        </div>
        <div class="field">
          <label>Phone Number</label>
          <input type="tel" id="sms-phone" placeholder="+254712345678">
        </div>
        <div class="field">
          <label>Message</label>
          <input type="text" id="sms-text" placeholder="Yes, let's talk.">
        </div>
        <button class="btn btn-primary" id="sms-btn" onclick="simulateSMS()">
          <span id="sms-spin" style="display:none" class="btn-spinner"></span>
          Send SMS
        </button>
      </div>
      <pre class="result-box" id="sms-result"></pre>
    </div>
  </div>

  <!-- ── Reply Simulation ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">💬</span>Reply Simulation</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <div class="form-row">
        <div class="field">
          <label>Lead Email</label>
          <input type="email" id="reply-email" placeholder="prospect@company.com">
        </div>
        <div class="field">
          <label>Reply Text</label>
          <input type="text" id="reply-text" placeholder="I'm interested, tell me more.">
        </div>
        <button class="btn btn-primary" id="reply-btn" onclick="simulateReply()">
          <span id="reply-spin" style="display:none" class="btn-spinner"></span>
          Send Reply
        </button>
      </div>
      <pre class="result-box" id="reply-result"></pre>
    </div>
  </div>

  <!-- ── Batch Upload & Run ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">📁</span>Batch Upload &amp; Run</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <p style="font-size:12px;color:var(--muted);margin-bottom:14px">
        Upload a CSV, PDF, or DOCX file with company info. Set how many companies to run,
        then start the batch pipeline (auto-sends emails using compose &amp; send, no per-email approval).
      </p>

      <!-- Upload zone -->
      <div class="upload-zone" id="upload-zone"
           ondragover="event.preventDefault();this.classList.add('drag-over')"
           ondragleave="this.classList.remove('drag-over')"
           ondrop="handleDrop(event)"
           onclick="document.getElementById('batch-file-input').click()">
        <div class="upload-icon">📂</div>
        <div class="upload-label">Drop file here or click to browse</div>
        <div class="upload-hint">Accepts .csv · .pdf · .docx</div>
        <input type="file" id="batch-file-input" accept=".csv,.pdf,.docx"
               style="display:none" onchange="handleFileSelect(event)">
      </div>
      <div id="batch-parse-error" style="display:none;color:var(--red);font-size:12px;margin-top:8px"></div>

      <!-- Preview + controls (shown after parse) -->
      <div id="batch-preview" style="display:none;margin-top:16px">
        <div class="batch-controls">
          <span class="batch-count-badge" id="batch-count-badge">0 companies</span>
          <div class="field">
            <label>Run first N companies</label>
            <input type="number" id="batch-n" min="1" value="5" style="width:90px">
          </div>
          <label class="checkbox-row">
            <input type="checkbox" id="batch-auto-approve" checked>
            Auto-send (skip per-email review)
          </label>
          <button class="btn btn-primary" id="batch-run-btn" onclick="startBatch()">
            <span id="batch-run-spin" style="display:none" class="btn-spinner"></span>
            ▶ Start Batch Run
          </button>
          <button class="btn btn-ghost btn-sm" onclick="clearBatch()">✕ Clear</button>
        </div>

        <!-- Parsed leads preview table -->
        <div class="tbl-wrap" style="max-height:240px;overflow-y:auto;margin-bottom:14px">
          <table class="batch-table">
            <thead>
              <tr>
                <th class="num-cell">#</th>
                <th>Email</th>
                <th>Company</th>
                <th>Domain</th>
                <th>Source</th>
              </tr>
            </thead>
            <tbody id="batch-preview-tbody"></tbody>
          </table>
        </div>
      </div>

      <!-- Batch run progress -->
      <div id="batch-progress-section" style="display:none;margin-top:14px">
        <div class="progress-wrap">
          <div class="progress-label">
            <span id="batch-status-text">Running…</span>
            <span id="batch-pct-text">0%</span>
          </div>
          <div class="progress-track">
            <div class="progress-fill" id="batch-fill" style="width:0%"></div>
          </div>
        </div>
        <div id="batch-items-list" style="margin-top:10px;max-height:260px;overflow-y:auto"></div>
      </div>
    </div>
  </div>

  <!-- ── Resend Setup Guide ── -->
  <div class="panel">
    <div class="panel-head" onclick="togglePanel(this)">
      <span class="panel-title"><span class="panel-icon">📬</span>Resend Inbound Email Setup Guide</span>
      <span class="chevron">▼</span>
    </div>
    <div class="panel-body">
      <div class="info-block">
        <h4>How to configure Resend to receive prospect replies</h4>
        <ol>
          <li><strong>Add your domain in Resend</strong> — go to <em>Resend Dashboard → Domains → Add Domain</em>. Use a subdomain like <code>mail.yourdomain.com</code> for sending and <code>in.yourdomain.com</code> for receiving.</li>
          <li><strong>Verify sending domain</strong> — add the DKIM + SPF DNS records Resend shows you. Wait for "Verified" status.</li>
          <li><strong>Enable inbound</strong> — in Resend Dashboard go to <em>Email → Inbound</em>, click <em>Add Inbound</em>, choose your receiving domain (e.g. <code>in.yourdomain.com</code>).</li>
          <li><strong>Add MX record to DNS</strong>:
            <br><code>Type: MX | Host: in | Value: inbound.resend.com | Priority: 10</code>
          </li>
          <li><strong>Set webhook URL</strong> — in the inbound route config, paste your app's inbound webhook URL:
            <br><code id="inbound-wh-url" style="user-select:all;background:var(--accent-light);padding:2px 6px;border-radius:4px;color:var(--accent)">loading…</code>
            <br><small style="color:var(--muted)">Update <code>WEBHOOK_BASE_URL</code> in <code>.env</code> when you deploy to production.</small>
          </li>
          <li><strong>Add webhook secret</strong> — copy the signing secret from Resend → set <code>RESEND_WEBHOOK_SECRET=whsec_…</code> in <code>.env</code>.
            <br><span id="wh-secret-status" style="font-size:12px"></span>
          </li>
          <li><strong>Set Reply-To on outbound emails</strong> — add <code>RESEND_REPLY_TO=replies@in.yourdomain.com</code> in <code>.env</code> so prospect replies route back through Resend inbound. (The app reads this env var automatically.)</li>
          <li><strong>Go live</strong> — set <code>OUTBOUND_LIVE=true</code> to route real emails to prospects (currently in sink mode — all mail goes to <code>STAFF_SINK_EMAIL</code> for safety).</li>
        </ol>
      </div>
    </div>
  </div>

  <!-- ── Leads Table ── -->
  <div style="display:flex;align-items:center;justify-content:space-between;margin-bottom:12px;margin-top:4px">
    <h2 style="font-size:15px;font-weight:700">Leads</h2>
  </div>
  <div class="toolbar">
    <input class="search-input" type="text" id="search"
           placeholder="Search by email, company…" oninput="filterTable()">
    <select id="status-filter" onchange="filterTable()">
      <option value="">All statuses</option>
      <option value="new">New</option>
      <option value="outreach_sent">Outreach Sent</option>
      <option value="in_conversation">In Conversation</option>
      <option value="qualified">Qualified</option>
      <option value="disqualified">Disqualified</option>
    </select>
  </div>
  <div class="panel" style="margin-bottom:0">
    <div class="tbl-wrap">
      <table>
        <thead>
          <tr>
            <th>Email</th><th>Company</th><th>Status</th>
            <th>Segment</th><th>AI Score</th><th>Turns</th><th>Created</th><th></th>
          </tr>
        </thead>
        <tbody id="leads-tbody">
          <tr><td colspan="8" class="empty-state">Loading…</td></tr>
        </tbody>
      </table>
    </div>
  </div>

</main>

<!-- ══════════════════ SIDEBAR ══════════════════ -->
<aside class="sidebar">
  <div class="sidebar-title">Pipeline Runs</div>
  <div id="run-history"><div class="sidebar-empty">No runs yet.<br>Start the pipeline above.</div></div>
</aside>

</div><!-- .layout -->

<!-- ── Email Approval Modal ── -->
<div class="overlay" id="approval-overlay">
  <div class="modal modal-lg">
    <div class="modal-head">
      <div>
        <h3>Review Email Before Sending</h3>
        <div class="sub">Edit subject or body, then approve to send — or reject to cancel the run.</div>
      </div>
      <button class="close-btn" onclick="rejectEmail()">✕</button>
    </div>
    <div class="modal-body">
      <div class="approval-hint">
        ✏️ You can freely edit the subject and body below before approving.
        The pipeline will only send after you click <strong>Approve &amp; Send</strong>.
      </div>
      <div class="approval-to-row">
        <span class="approval-to-label">To</span>
        <span class="approval-to-email" id="ap-to"></span>
        <span id="ap-sink-note" class="sink-note" style="display:none">→ redirected to sink</span>
      </div>
      <div class="field" style="margin-bottom:10px">
        <label>Subject</label>
        <input type="text" id="ap-subject">
      </div>
      <div class="field">
        <label>Body</label>
        <textarea id="ap-body" style="min-height:200px"></textarea>
      </div>
    </div>
    <div class="modal-footer">
      <button class="btn btn-danger" onclick="rejectEmail()">✗ Reject & Cancel Run</button>
      <button class="btn btn-success" onclick="approveEmail()">
        <span id="ap-send-spin" style="display:none" class="btn-spinner"></span>
        ✓ Approve &amp; Send
      </button>
    </div>
  </div>
</div>

<!-- ── Lead Detail Modal ── -->
<div class="overlay" id="lead-modal" onclick="closeLeadModal(event)">
  <div class="modal">
    <div class="modal-head">
      <div>
        <h3 id="modal-title">Lead Detail</h3>
        <div class="sub" id="modal-sub"></div>
      </div>
      <button class="close-btn" onclick="closeLeadModalDirect()">✕</button>
    </div>
    <div class="modal-body" id="modal-body"></div>
  </div>
</div>

<script>
'use strict';

// ── State ─────────────────────────────────────────────────────────────────────
let allLeads       = [];
let currentRunId   = null;
let pollTimer      = null;
let approvalShown  = false;
const TOTAL_STEPS  = 6;
const ALL_STEPS    = ['Enrichment','Signals Research','Competitor Gap','Email Composition','Email Send','HubSpot Sync'];

// ── Init ──────────────────────────────────────────────────────────────────────
// Apply server-injected config (SVRCFG is injected by the Python route handler)
(function applyServerConfig() {
  if (typeof SVRCFG === 'undefined') return;
  // Outbound badge
  const badge = document.getElementById('outbound-badge');
  if (badge) {
    if (SVRCFG.outboundLive) {
      badge.className = 'outbound-badge outbound-live';
      badge.textContent = 'LIVE MODE';
    } else {
      badge.className = 'outbound-badge outbound-sink';
      badge.textContent = `SINK → ${SVRCFG.staffSink}`;
    }
  }
  // Resend guide webhook URL + secret status
  const whEl = document.getElementById('inbound-wh-url');
  if (whEl) whEl.textContent = SVRCFG.webhookUrl;
  const secEl = document.getElementById('wh-secret-status');
  if (secEl) {
    secEl.textContent = SVRCFG.secretSet
      ? '✓ RESEND_WEBHOOK_SECRET is set — signatures will be verified.'
      : '✗ RESEND_WEBHOOK_SECRET not set — all webhook requests are accepted (dev mode).';
    secEl.style.color = SVRCFG.secretSet ? 'var(--green)' : 'var(--yellow)';
  }
})();

checkHealth();
loadLeads();
refreshHistory();
setInterval(checkHealth, 30000);
setInterval(refreshHistory, 5000);

// ── Health ────────────────────────────────────────────────────────────────────
async function checkHealth() {
  try {
    const r = await fetch('/health');
    const d = await r.json();
    if (d.status === 'ok') {
      document.getElementById('hDot').classList.add('ok');
      document.getElementById('hLabel').textContent = 'Live';
    }
  } catch(_) {}
}

// ── Leads ─────────────────────────────────────────────────────────────────────
async function loadLeads() {
  const res = await fetch('/api/leads');
  allLeads  = await res.json();
  updateStats(allLeads);
  renderTable(allLeads);
}

function updateStats(leads) {
  const cnt = {new:0, outreach_sent:0, in_conversation:0, qualified:0, disqualified:0};
  leads.forEach(l => { cnt[l.status] = (cnt[l.status]||0) + 1; });
  document.getElementById('s-total').textContent = leads.length;
  document.getElementById('s-new').textContent   = cnt.new;
  document.getElementById('s-sent').textContent  = cnt.outreach_sent;
  document.getElementById('s-conv').textContent  = cnt.in_conversation;
  document.getElementById('s-qual').textContent  = cnt.qualified;
  document.getElementById('s-disq').textContent  = cnt.disqualified;
}

function statusBadge(s) {
  const map = {
    new:             ['badge-new',  'New'],
    outreach_sent:   ['badge-sent', 'Sent'],
    in_conversation: ['badge-conv', 'In Convo'],
    qualified:       ['badge-qual', 'Qualified'],
    disqualified:    ['badge-disq', 'Disqualified'],
  };
  const [cls, lbl] = map[s] || ['badge-new', s];
  return `<span class="badge ${cls}">${lbl}</span>`;
}

function scoreDot(n) {
  const cls = ['s0','s1','s2','s3','s4'][Math.min(n||0, 4)];
  return `<span class="score"><span class="score-dot ${cls}"></span>${n||0}</span>`;
}

function fmtDate(s) {
  if (!s) return '—';
  return s.replace('T',' ').replace('Z','').slice(0,16);
}

function b64(email) { return btoa(email).replace(/=/g,''); }

function renderTable(leads) {
  const tbody = document.getElementById('leads-tbody');
  if (!leads.length) {
    tbody.innerHTML = '<tr><td colspan="8" class="empty-state">No leads yet. Run the pipeline to get started.</td></tr>';
    return;
  }
  tbody.innerHTML = leads.map(l => `
    <tr onclick="openLeadModal('${l.email.replace(/'/g,"\\'")}')">
      <td class="email-cell">${l.email}</td>
      <td class="company-cell">${l.company_name || '<span style="color:var(--muted)">—</span>'}</td>
      <td>${statusBadge(l.status)}</td>
      <td style="font-size:12px;color:var(--muted)">${(l.segment_label||'generic').replace(/_/g,' ')}</td>
      <td>${scoreDot(l.ai_maturity_score)}</td>
      <td style="color:var(--muted)">${l.turns}</td>
      <td style="font-size:12px;color:var(--muted)">${fmtDate(l.created_at)}</td>
      <td onclick="event.stopPropagation()">
        <button class="btn btn-danger btn-sm" onclick="deleteLead('${l.email.replace(/'/g,"\\'")}')">✕</button>
      </td>
    </tr>`).join('');
}

function filterTable() {
  const q  = document.getElementById('search').value.toLowerCase();
  const sf = document.getElementById('status-filter').value;
  renderTable(allLeads.filter(l => {
    const mQ  = !q  || l.email.toLowerCase().includes(q) || (l.company_name||'').toLowerCase().includes(q);
    const mSF = !sf || l.status === sf;
    return mQ && mSF;
  }));
}

async function deleteLead(email) {
  if (!confirm(`Delete lead ${email}?`)) return;
  await fetch(`/api/leads/${b64(email)}`, {method:'DELETE'});
  await loadLeads();
}

// ── Lead Modal ────────────────────────────────────────────────────────────────
async function openLeadModal(email) {
  const res = await fetch(`/api/leads/${b64(email)}`);
  const l   = await res.json();
  document.getElementById('modal-title').textContent = l.company_name || email;
  document.getElementById('modal-sub').textContent   = email;
  const p = l.profile || {};
  const rows = [
    ['Status',      statusBadge(l.status)],
    ['Lead ID',     `<code style="font-size:11px">${l.lead_id}</code>`],
    ['Phone',       l.phone || '—'],
    ['HubSpot ID',  l.hubspot_contact_id || '—'],
    ['Domain',      l.domain || '—'],
    ['Segment',     (l.segment_label||'generic').replace(/_/g,' ')],
    ['AI Maturity', scoreDot(l.ai_maturity_score)],
    ['Turns',       l.turns],
    ['Enriched',    fmtDate(l.enriched_at)],
    ['Created',     fmtDate(l.created_at)],
    ['Booking',     l.booking_url ? `<a href="${l.booking_url}" target="_blank" style="color:var(--accent)">${l.booking_url}</a>` : '—'],
  ];
  let html = `<div class="kv-grid">${rows.map(([k,v])=>`<div class="kv"><div class="k">${k}</div><div class="v">${v}</div></div>`).join('')}</div>`;
  if (l.history && l.history.length) {
    html += `<div class="section-label">Conversation (${l.history.length} turns)</div>`;
    html += l.history.map(m=>`<div class="history-item"><div class="role role-${m.role}">${m.role}</div><div class="msg">${m.content||m.text||JSON.stringify(m)}</div></div>`).join('');
  }
  if (p.competitor_gap_brief) {
    html += `<div class="section-label">Competitor Gap Brief</div><div class="json-block">${p.competitor_gap_brief}</div>`;
  }
  if (p.signals_research) {
    html += `<div class="section-label">Signals Research</div><div class="json-block">${JSON.stringify(p.signals_research,null,2)}</div>`;
  }
  document.getElementById('modal-body').innerHTML = html;
  document.getElementById('lead-modal').classList.add('open');
}
function closeLeadModal(e) { if (e.target.id==='lead-modal') closeLeadModalDirect(); }
function closeLeadModalDirect() { document.getElementById('lead-modal').classList.remove('open'); }

// ── Panel toggle ──────────────────────────────────────────────────────────────
function togglePanel(head) {
  head.classList.toggle('open');
  head.nextElementSibling.classList.toggle('open');
}

// ── Pipeline ──────────────────────────────────────────────────────────────────
async function startPipeline() {
  const email = document.getElementById('pipe-email').value.trim();
  const text  = document.getElementById('pipe-text').value.trim();
  if (!email) { alert('Lead email is required'); return; }
  if (pollTimer) { clearInterval(pollTimer); pollTimer = null; }
  approvalShown = false;

  // Reset UI
  document.getElementById('pipe-btn').disabled = true;
  document.getElementById('pipe-spin').style.display = 'inline-block';
  document.getElementById('pipe-progress-wrap').style.display = 'block';
  document.getElementById('pipe-fill').style.width = '0%';
  document.getElementById('pipe-pct').textContent = '0%';
  document.getElementById('pipe-status-text').textContent = 'Starting…';
  document.getElementById('pipe-msg').textContent = '';
  renderSteps({}, []);

  const res = await fetch('/api/pipeline/run', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({email, text: text||'Tell me more.'})
  });
  const {run_id} = await res.json();
  currentRunId = run_id;
  pollTimer = setInterval(()=>pollRun(run_id), 1000);
}

async function pollRun(run_id) {
  try {
    const res = await fetch(`/api/pipeline/${run_id}`);
    if (!res.ok) return;
    const run = await res.json();
    updatePipelineUI(run);

    if (run.status === 'awaiting_approval' && !approvalShown) {
      approvalShown = true;
      showApprovalModal(run.email_draft);
    }

    if (['completed','failed','rejected'].includes(run.status)) {
      clearInterval(pollTimer); pollTimer = null;
      document.getElementById('pipe-btn').disabled = false;
      document.getElementById('pipe-spin').style.display = 'none';
      refreshHistory();
      loadLeads();
    }
  } catch(e) { console.error('poll error', e); }
}

function updatePipelineUI(run) {
  const steps = run.steps || {};
  let doneW = 0;
  ALL_STEPS.forEach(n => {
    const s = steps[n];
    if (!s) return;
    if (s.status === 'done') doneW += 1;
    else if (s.status === 'running') doneW += 0.5;
    else if (s.status === 'awaiting_approval') doneW += 0.7;
    else if (s.status === 'error' || s.status === 'rejected') doneW += 1;
  });
  const pct = Math.min(Math.round((doneW / TOTAL_STEPS) * 100), 100);
  document.getElementById('pipe-fill').style.width = pct + '%';
  document.getElementById('pipe-pct').textContent  = pct + '%';

  const msgMap = {
    completed:         '✓ Pipeline complete!',
    failed:            `✗ Failed: ${run.error||'unknown error'}`,
    rejected:          '✗ Email rejected — run cancelled.',
    awaiting_approval: '⏸ Waiting for your email approval…',
    approved:          '↻ Sending email…',
    running:           '↻ Running…',
    pending:           'Starting…',
  };
  document.getElementById('pipe-status-text').textContent = msgMap[run.status] || run.status;
  renderSteps(steps, run.step_order || []);
}

function renderSteps(steps, stepOrder) {
  const iconMap = {
    pending:           ['○', '—'],
    running:           [`<span class="spin">↻</span>`, 'Running'],
    done:              ['✓', 'Done'],
    error:             ['✗', 'Error'],
    awaiting_approval: ['⏸', 'Review'],
    rejected:          ['✗', 'Rejected'],
  };
  document.getElementById('pipe-steps').innerHTML = ALL_STEPS.map(name => {
    const step   = steps[name];
    const status = step ? step.status : 'pending';
    const [icon, badge] = iconMap[status] || ['○', status];
    let detail = '';
    if (step && step.data) {
      const vals = Object.values(step.data);
      if (vals.length) detail = String(vals[0]).slice(0, 70);
    }
    if (step && step.error) detail = step.error.slice(0, 70);
    return `<div class="step-card s-${status}">
      <div class="step-icon-wrap s-${status}">${icon}</div>
      <div class="step-content">
        <div class="step-name">${name}</div>
        ${detail ? `<div class="step-detail">${detail}</div>` : ''}
      </div>
      <span class="step-badge s-${status}">${badge}</span>
    </div>`;
  }).join('');
}

// ── Email Approval ────────────────────────────────────────────────────────────
function showApprovalModal(draft) {
  if (!draft) return;
  window._pendingDraft = draft;
  document.getElementById('ap-to').textContent      = draft.to;
  document.getElementById('ap-subject').value       = draft.subject || '';
  document.getElementById('ap-body').value          = draft.body   || '';
  const sinkNote = document.getElementById('ap-sink-note');
  const isLive = typeof SVRCFG !== 'undefined' ? SVRCFG.outboundLive
               : document.getElementById('outbound-badge').classList.contains('outbound-live');
  sinkNote.style.display = isLive ? 'none' : 'inline-flex';
  document.getElementById('approval-banner').classList.add('show');
  document.getElementById('approval-overlay').classList.add('open');
  // Scroll the modal into view in case user has scrolled down
  document.getElementById('approval-overlay').scrollIntoView({behavior:'smooth', block:'center'});
}

async function approveEmail() {
  if (!currentRunId) return;
  const subject = document.getElementById('ap-subject').value.trim();
  const body    = document.getElementById('ap-body').value.trim();
  if (!subject || !body) { alert('Subject and body cannot be empty.'); return; }
  document.getElementById('ap-send-spin').style.display = 'inline-block';
  try {
    const res = await fetch(`/api/pipeline/${currentRunId}/approve`, {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({subject, body})
    });
    if (!res.ok) {
      const err = await res.json().catch(()=>({error:String(res.status)}));
      alert(`Approval failed: ${err.error || res.status}`);
      document.getElementById('ap-send-spin').style.display = 'none';
      return;
    }
  } catch(e) {
    alert(`Network error: ${e.message}`);
    document.getElementById('ap-send-spin').style.display = 'none';
    return;
  }
  document.getElementById('ap-send-spin').style.display = 'none';
  document.getElementById('approval-banner').classList.remove('show');
  document.getElementById('approval-overlay').classList.remove('open');
  approvalShown = false;
  window._pendingDraft = null;
}

async function rejectEmail() {
  if (currentRunId) {
    if (!confirm('Reject this email? The pipeline run will be cancelled.')) return;
    await fetch(`/api/pipeline/${currentRunId}/reject`, {method:'POST'});
  }
  document.getElementById('approval-banner').classList.remove('show');
  document.getElementById('approval-overlay').classList.remove('open');
  approvalShown = false;
  window._pendingDraft = null;
}

// ── History ───────────────────────────────────────────────────────────────────
async function refreshHistory() {
  const res = await fetch('/api/pipeline/history');
  const hist = await res.json();
  renderHistory(hist);
  // Check live badge
  const res2 = await fetch('/health').catch(()=>null);
  // outbound-live badge is set server-side on page load via the HTML attribute
}

function renderHistory(history) {
  const el = document.getElementById('run-history');
  if (!history.length) {
    el.innerHTML = '<div class="sidebar-empty">No runs yet.<br>Start the pipeline above.</div>';
    return;
  }
  el.innerHTML = history.map(run => {
    const statusPill = {
      completed: 'rp-completed', failed: 'rp-failed',
      rejected:  'rp-rejected',  running: 'rp-running',
      awaiting_approval: 'rp-awaiting', pending: 'rp-pending',
    }[run.status] || 'rp-pending';
    const statusLabel = {
      completed: '✓ Done', failed: '✗ Failed',
      rejected:  'Rejected', running: '↻ Running',
      awaiting_approval: '⏸ Waiting', pending: 'Pending',
    }[run.status] || run.status;
    const elapsed = run.completed_at && run.started_at
      ? ((run.completed_at - run.started_at).toFixed(1) + 's')
      : 'in progress';
    const ts = run.started_at
      ? new Date(run.started_at * 1000).toLocaleTimeString([], {hour:'2-digit',minute:'2-digit'})
      : '';
    let detailHtml = '';
    if (run.error) detailHtml += `<div class="run-error-msg">Error: ${run.error}</div>`;
    if (run.result) {
      const r = run.result;
      detailHtml += `<pre class="run-result-pre">${JSON.stringify(r,null,2).slice(0,300)}</pre>`;
    }
    if (run.gap_brief) {
      detailHtml += `<div class="run-gap-preview">${run.gap_brief.slice(0,180)}…</div>`;
    }
    return `<div class="run-card">
      <div class="run-card-head" onclick="toggleRunDetail('${run.run_id}')">
        <div class="run-email">${run.email}</div>
        <div class="run-meta">
          <span class="run-time">${ts} · ${elapsed}</span>
          <span class="run-status-pill ${statusPill}">${statusLabel}</span>
        </div>
      </div>
      ${detailHtml ? `<div class="run-detail" id="rd-${run.run_id}">${detailHtml}</div>` : ''}
    </div>`;
  }).join('');
}

function toggleRunDetail(runId) {
  const el = document.getElementById(`rd-${runId}`);
  if (el) el.classList.toggle('open');
}

// ── Email Outreach Panel ──────────────────────────────────────────────────────
let emTraceId = '';
async function composeEmail() {
  const email = document.getElementById('em-email').value.trim();
  if (!email) { alert('Lead email required'); return; }
  setBtnLoading('em-compose-btn','em-compose-spin',true);
  document.getElementById('em-draft-wrap').style.display = 'none';
  document.getElementById('em-result').className = 'result-box';
  try {
    const res  = await fetch('/api/email/compose', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({email})
    });
    const data = await res.json();
    if (!res.ok) throw new Error(data.error || 'Compose failed');
    document.getElementById('em-to').textContent      = data.to;
    document.getElementById('em-subject').value       = data.subject;
    document.getElementById('em-body').value          = data.body;
    emTraceId = data.trace_id;
    const isLive = typeof SVRCFG !== 'undefined' ? SVRCFG.outboundLive
                 : document.getElementById('outbound-badge').classList.contains('outbound-live');
    document.getElementById('em-sink-note').style.display = isLive ? 'none' : 'inline-flex';
    document.getElementById('em-draft-wrap').style.display = 'block';
  } catch(e) {
    showResult('em-result', String(e), true);
  } finally {
    setBtnLoading('em-compose-btn','em-compose-spin',false);
  }
}

async function sendComposedEmail() {
  const to      = document.getElementById('em-to').textContent.trim();
  const subject = document.getElementById('em-subject').value.trim();
  const body    = document.getElementById('em-body').value.trim();
  if (!to || !subject || !body) { alert('Subject and body required'); return; }
  setBtnLoading(null,'em-send-spin',true);
  try {
    const res  = await fetch('/api/email/send', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({to, subject, body, trace_id: emTraceId})
    });
    const data = await res.json();
    showResult('em-result', JSON.stringify(data,null,2), !res.ok);
    if (res.ok) {
      document.getElementById('em-draft-wrap').style.display = 'none';
      loadLeads();
    }
  } catch(e) {
    showResult('em-result', String(e), true);
  } finally {
    setBtnLoading(null,'em-send-spin',false);
  }
}

// ── Gap Analysis Panel ────────────────────────────────────────────────────────
async function analyzeGap() {
  const email = document.getElementById('gap-email').value.trim();
  if (!email) { alert('Lead email required'); return; }
  await callPanelApi('/api/gap/analyze', {email}, 'gap-btn', 'gap-spin', 'gap-result');
}

// ── CRM Panel ─────────────────────────────────────────────────────────────────
async function syncCRM() {
  const email = document.getElementById('crm-email').value.trim();
  if (!email) { alert('Lead email required'); return; }
  await callPanelApi('/api/crm/sync', {email}, 'crm-btn', 'crm-spin', 'crm-result');
}

// ── SMS Panel ─────────────────────────────────────────────────────────────────
async function simulateSMS() {
  const email = document.getElementById('sms-email').value.trim();
  const phone = document.getElementById('sms-phone').value.trim();
  const text  = document.getElementById('sms-text').value.trim();
  if (!email || !phone) { alert('Email and phone required'); return; }
  await callPanelApi('/simulate/sms', {email, phone, text: text||"Yes, let's talk."}, 'sms-btn','sms-spin','sms-result');
  loadLeads();
}

// ── Reply Panel ───────────────────────────────────────────────────────────────
async function simulateReply() {
  const email = document.getElementById('reply-email').value.trim();
  const text  = document.getElementById('reply-text').value.trim();
  if (!email || !text) { alert('Email and text required'); return; }
  await callPanelApi('/webhooks/email', {from:email, text, thread_id:'dashboard-reply'}, 'reply-btn','reply-spin','reply-result');
  loadLeads();
}

// ── Helpers ───────────────────────────────────────────────────────────────────
function showResult(id, text, isError) {
  const el = document.getElementById(id);
  el.textContent  = text;
  el.className    = 'result-box visible' + (isError ? ' error' : '');
}

function setBtnLoading(btnId, spinId, on) {
  if (btnId) document.getElementById(btnId).disabled = on;
  if (spinId) document.getElementById(spinId).style.display = on ? 'inline-block' : 'none';
}

async function callPanelApi(endpoint, payload, btnId, spinId, resultId) {
  setBtnLoading(btnId, spinId, true);
  document.getElementById(resultId).className = 'result-box';
  try {
    const res  = await fetch(endpoint, {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify(payload)
    });
    const data = await res.json();
    showResult(resultId, JSON.stringify(data,null,2), !res.ok);
  } catch(e) {
    showResult(resultId, String(e), true);
  } finally {
    setBtnLoading(btnId, spinId, false);
  }
}

// ── Batch Upload ──────────────────────────────────────────────────────────────
let batchLeads   = [];
let batchPollTmr = null;

function handleDrop(e) {
  e.preventDefault();
  document.getElementById('upload-zone').classList.remove('drag-over');
  const file = e.dataTransfer.files[0];
  if (file) uploadBatchFile(file);
}
function handleFileSelect(e) {
  const file = e.target.files[0];
  if (file) uploadBatchFile(file);
  e.target.value = '';
}

async function uploadBatchFile(file) {
  const errEl = document.getElementById('batch-parse-error');
  errEl.style.display = 'none';
  document.getElementById('batch-preview').style.display = 'none';
  document.getElementById('batch-progress-section').style.display = 'none';
  const zone = document.getElementById('upload-zone');
  zone.querySelector('.upload-label').textContent = `Parsing ${file.name}…`;

  const fd = new FormData();
  fd.append('file', file);
  try {
    const res  = await fetch('/api/batch/parse', {method:'POST', body: fd});
    const data = await res.json();
    if (!res.ok) throw new Error(data.error || 'Parse failed');
    batchLeads = data.leads;
    zone.querySelector('.upload-label').textContent = `${file.name} — ${data.count} entries found`;
    renderBatchPreview(data.leads);
  } catch(e) {
    zone.querySelector('.upload-label').textContent = 'Drop file here or click to browse';
    errEl.textContent = String(e);
    errEl.style.display = 'block';
  }
}

function renderBatchPreview(leads) {
  document.getElementById('batch-count-badge').textContent = `${leads.length} companies`;
  document.getElementById('batch-n').max   = leads.length;
  document.getElementById('batch-n').value = Math.min(leads.length, 5);
  const tbody = document.getElementById('batch-preview-tbody');
  tbody.innerHTML = leads.map((l,i) => `
    <tr>
      <td class="num-cell">${i+1}</td>
      <td class="email-cell">${l.email}</td>
      <td>${l.company || '<span style="color:var(--muted)">—</span>'}</td>
      <td style="font-size:12px;color:var(--muted)">${l.domain}</td>
      <td><span class="badge ${l.source==='direct'?'badge-qual':l.source==='synthesized'?'badge-conv':'badge-sent'}">${l.source}</span></td>
    </tr>`).join('');
  document.getElementById('batch-preview').style.display = 'block';
}

async function startBatch() {
  if (!batchLeads.length) { alert('No leads loaded'); return; }
  const n          = parseInt(document.getElementById('batch-n').value) || 1;
  const autoApprove = document.getElementById('batch-auto-approve').checked;
  if (!autoApprove) {
    alert('Manual approval for batch is not yet supported. Please enable "Auto-send" for batch runs.');
    return;
  }
  setBtnLoading('batch-run-btn','batch-run-spin',true);
  document.getElementById('batch-progress-section').style.display = 'block';
  document.getElementById('batch-items-list').innerHTML = '';
  document.getElementById('batch-fill').style.width = '0%';
  document.getElementById('batch-pct-text').textContent = '0%';
  document.getElementById('batch-status-text').textContent = 'Starting…';

  const res = await fetch('/api/batch/run', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({leads: batchLeads, n})
  });
  const data = await res.json();
  if (!res.ok) { alert(data.error); setBtnLoading('batch-run-btn','batch-run-spin',false); return; }

  batchPollTmr = setInterval(pollBatch, 1500);
}

async function pollBatch() {
  try {
    const res  = await fetch('/api/batch/status');
    const data = await res.json();
    updateBatchUI(data);
    if (data.status === 'completed' || data.status === 'failed') {
      clearInterval(batchPollTmr); batchPollTmr = null;
      setBtnLoading('batch-run-btn','batch-run-spin',false);
      loadLeads();
      refreshHistory();
    }
  } catch(e) { console.error('batch poll error', e); }
}

function updateBatchUI(data) {
  const total = data.total || 1;
  const done  = data.done + data.failed_count;
  const pct   = Math.min(Math.round((done / total) * 100), 100);
  document.getElementById('batch-fill').style.width = pct + '%';
  document.getElementById('batch-pct-text').textContent = pct + '%';

  const statusMap = {
    running:   `↻ Running — ${data.done}/${total} done, ${data.failed_count} failed` + (data.current ? ` | current: ${data.current}` : ''),
    completed: `✓ Batch complete — ${data.done} done, ${data.failed_count} failed`,
    failed:    `✗ Batch failed`,
    idle:      'Idle',
  };
  document.getElementById('batch-status-text').textContent = statusMap[data.status] || data.status;

  const listEl = document.getElementById('batch-items-list');
  listEl.innerHTML = (data.latest_results || []).slice().reverse().map(r => {
    const isOk = r.status === 'ok';
    return `<div class="batch-item">
      <span class="bi-email">${r.email}</span>
      <span class="bi-status ${isOk?'bi-ok':'bi-error'}">${isOk?'✓ Sent':'✗ Error'}</span>
      ${!isOk ? `<span style="font-size:11px;color:var(--red)">${(r.error||'').slice(0,60)}</span>` : ''}
    </div>`;
  }).join('') + (data.current && data.status === 'running' ? `
    <div class="batch-item">
      <span class="bi-email">${data.current}</span>
      <span class="bi-status bi-running"><span class="spin">↻</span> Running</span>
    </div>` : '');
}

function clearBatch() {
  batchLeads = [];
  document.getElementById('batch-preview').style.display = 'none';
  document.getElementById('batch-progress-section').style.display = 'none';
  document.getElementById('upload-zone').querySelector('.upload-label').textContent = 'Drop file here or click to browse';
  document.getElementById('batch-parse-error').style.display = 'none';
}

// ── Outbound badge ────────────────────────────────────────────────────────────
(async function detectOutboundMode() {
  try {
    await fetch('/health');
    // Badge stays "SINK MODE" by default — set OUTBOUND_LIVE=true in .env to switch
  } catch(_){}
})();
</script>
</body>
</html>"""


@app.get("/", response_class=HTMLResponse)
async def dashboard():
    base = os.getenv("WEBHOOK_BASE_URL", "https://your-app.onrender.com").rstrip("/")
    webhook_url = f"{base}/webhooks/email"
    secret_set = "true" if os.getenv("RESEND_WEBHOOK_SECRET") else "false"
    live = "true" if _OUTBOUND_LIVE else "false"
    sink = _STAFF_SINK.replace('"', r'\"')
    wh_esc = webhook_url.replace('"', r'\"')
    config_script = (
        f'<script>const SVRCFG={{'
        f'"webhookUrl":"{wh_esc}",'
        f'"outboundLive":{live},'
        f'"staffSink":"{sink}",'
        f'"secretSet":{secret_set}'
        f'}};</script>'
    )
    return HTMLResponse(_HTML.replace("<!-- __SERVER_CONFIG__ -->", config_script))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("dashboard:app", host="0.0.0.0", port=8080, reload=True)
